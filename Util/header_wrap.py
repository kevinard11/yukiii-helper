from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


EMU_PER_PT = 12700


def add_header_image_and_wrap(header, image_path, pos_x_pt=0, pos_y_pt=0):
    """
    Safe floating header image without breaking Microsoft Word.
    """
    paragraph = header.add_paragraph()
    run = paragraph.add_run()
    pic = run.add_picture(image_path)

    # Ambil inline (aman)
    inline = pic._inline

    # --- copy internal properties ---
    extent = inline.xpath('./wp:extent')[0]
    cx = extent.get("cx")
    cy = extent.get("cy")

    docPr = inline.xpath('./wp:docPr')[0]
    graphic = inline.xpath('./a:graphic')[0]

    # --- root anchor element ---
    anchor = OxmlElement("wp:anchor")
    anchor.set("behindDoc", "0")
    anchor.set("distT", "0")
    anchor.set("distB", "0")
    anchor.set("distL", "0")
    anchor.set("distR", "0")
    anchor.set("simplePos", "0")
    anchor.set("relativeHeight", "251658240")
    anchor.set("locked", "0")
    anchor.set("layoutInCell", "1")
    anchor.set("allowOverlap", "1")

    # --- simplePos ---
    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")
    anchor.append(simple_pos)

    # --- position (manual points → EMU) ---
    positionH = OxmlElement("wp:positionH")
    positionH.set("relativeFrom", "page")
    posH = OxmlElement("wp:posOffset")
    posH.text = str(int(pos_x_pt * EMU_PER_PT))
    positionH.append(posH)

    positionV = OxmlElement("wp:positionV")
    positionV.set("relativeFrom", "page")
    posV = OxmlElement("wp:posOffset")
    posV.text = str(int(pos_y_pt * EMU_PER_PT))
    positionV.append(posV)

    anchor.append(positionH)
    anchor.append(positionV)

    # --- size ---
    new_extent = OxmlElement("wp:extent")
    new_extent.set("cx", cx)
    new_extent.set("cy", cy)
    anchor.append(new_extent)

    # --- wrap square ---
    wrap_square = OxmlElement("wp:wrapSquare")
    wrap_square.set("wrapText", "bothSides")
    anchor.append(wrap_square)

    # --- append image XML ---
    anchor.append(docPr)
    anchor.append(graphic)

    # --- remove inline safely ---
    parent = inline.getparent()
    parent.remove(inline)

    # --- insert anchor ---
    parent.append(anchor)

    return header
